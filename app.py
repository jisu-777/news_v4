import streamlit as st
import re
from typing import Dict, List


# ✅ 무조건 첫 Streamlit 명령어
st.set_page_config(
    page_title="PwC 뉴스 분석기",
    page_icon="logo_orange.png",
    layout="wide",
)



from datetime import datetime, timedelta, timezone
import os
from PIL import Image
import docx
from docx.shared import Pt, RGBColor, Inches
import io
from urllib.parse import urlparse
from news_service import NewsAnalysisService
import pandas as pd  # 엑셀 생성을 위해 pandas 추가
import html  # HTML 엔티티 디코딩을 위해 추가

# Import centralized configuration
from config import (
    COMPANY_CATEGORIES,
    COMPANY_KEYWORD_MAP,
    COMPANY_GROUP_MAPPING,
    TRUSTED_PRESS_ALIASES,
    ADDITIONAL_PRESS_ALIASES,
    SYSTEM_PROMPT_1,
    SYSTEM_PROMPT_2,
    SYSTEM_PROMPT_3,
    EXCLUSION_CRITERIA,
    DUPLICATE_HANDLING,
    SELECTION_CRITERIA, 
    GPT_MODELS,
    DEFAULT_GPT_MODEL,
    # 새로 추가되는 회사별 기준들
    COMPANY_ADDITIONAL_EXCLUSION_CRITERIA,
    COMPANY_ADDITIONAL_DUPLICATE_HANDLING,
    COMPANY_ADDITIONAL_SELECTION_CRITERIA
)

# 한국 시간대(KST) 정의
KST = timezone(timedelta(hours=9))

def clean_html_entities(text):
    """HTML 엔티티를 정리하고 &quot; 등의 문제를 해결하는 함수"""
    if not text:
        return ""
    
    # HTML 엔티티 디코딩
    cleaned_text = html.unescape(str(text))
    
    # 추가적인 정리 작업
    cleaned_text = cleaned_text.replace('&quot;', '"')
    cleaned_text = cleaned_text.replace('&amp;', '&')
    cleaned_text = cleaned_text.replace('&lt;', '<')
    cleaned_text = cleaned_text.replace('&gt;', '>')
    cleaned_text = cleaned_text.replace('&apos;', "'")
    
    # 연속된 공백 정리
    cleaned_text = re.sub(r'\s+', ' ', cleaned_text).strip()
    
    return cleaned_text


def parse_press_config(press_dict_str: str) -> Dict[str, List[str]]:
    """UI에서 설정한 언론사 문자열을 딕셔너리로 파싱하는 함수"""
    press_config = {}
    if isinstance(press_dict_str, str) and press_dict_str.strip():
        try:
            lines = press_dict_str.strip().split('\n')
            for line in lines:
                line = line.strip()
                if line and ': ' in line:
                    press_name, aliases_str = line.split(':', 1)
                    try:
                        # 문자열 형태의 리스트를 실제 리스트로 변환
                        aliases = eval(aliases_str.strip())
                        press_config[press_name.strip()] = aliases
                    except Exception as e:
                        print(f"언론사 파싱 실패: {line}, 오류: {str(e)}")
        except Exception as e:
            print(f"전체 언론사 파싱 실패: {str(e)}")
    
    return press_config


def format_date(date_str):
    """Format date to MM/DD format with proper timezone handling"""
    try:
        # Try YYYY-MM-DD format
        date_obj = datetime.strptime(date_str, '%Y-%m-%d')
        return date_obj.strftime('%m/%d')
    except Exception:
        try:
            # Try GMT format and convert to KST
            date_obj = datetime.strptime(date_str, '%a, %d %b %Y %H:%M:%S %Z')
            # Convert UTC to KST (add 9 hours)
            date_obj_kst = date_obj + timedelta(hours=9)
            return date_obj_kst.strftime('%m/%d')
        except Exception:
            try:
                # Try GMT format without timezone indicator
                date_obj = datetime.strptime(date_str, '%a, %d %b %Y %H:%M:%S GMT')
                # Convert UTC to KST (add 9 hours)
                date_obj_kst = date_obj + timedelta(hours=9)
                return date_obj_kst.strftime('%m/%d')
            except Exception:
                # Return original if parsing fails
                return date_str if date_str else '날짜 정보 없음'

# 회사별 추가 기준을 적용하는 함수들
def get_enhanced_exclusion_criteria(companies):
    """회사별 제외 기준을 추가한 프롬프트 반환 (여러 회사 지원)"""
    base_criteria = EXCLUSION_CRITERIA
    
    # companies가 문자열이면 리스트로 변환
    if isinstance(companies, str):
        companies = [companies]
    
    # 선택된 모든 회사의 추가 기준을 합침
    all_additional_criteria = ""
    for company in companies:
        additional_criteria = COMPANY_ADDITIONAL_EXCLUSION_CRITERIA.get(company, "")
        if additional_criteria:
            all_additional_criteria += additional_criteria
    
    return base_criteria + all_additional_criteria

def get_enhanced_duplicate_handling(companies):
    """회사별 중복 처리 기준을 추가한 프롬프트 반환 (여러 회사 지원)"""
    base_criteria = DUPLICATE_HANDLING
    
    # companies가 문자열이면 리스트로 변환
    if isinstance(companies, str):
        companies = [companies]
    
    # 선택된 모든 회사의 추가 기준을 합침
    all_additional_criteria = ""
    for company in companies:
        additional_criteria = COMPANY_ADDITIONAL_DUPLICATE_HANDLING.get(company, "")
        if additional_criteria:
            all_additional_criteria += additional_criteria
    
    return base_criteria + all_additional_criteria

def get_enhanced_selection_criteria(companies):
    """회사별 선택 기준을 추가한 프롬프트 반환 (여러 회사 지원)"""
    base_criteria = SELECTION_CRITERIA
    
    # companies가 문자열이면 리스트로 변환
    if isinstance(companies, str):
        companies = [companies]
    
    # 선택된 모든 회사의 추가 기준을 합침
    all_additional_criteria = ""
    for company in companies:
        additional_criteria = COMPANY_ADDITIONAL_SELECTION_CRITERIA.get(company, "")
        if additional_criteria:
            all_additional_criteria += additional_criteria
    
    return base_criteria + all_additional_criteria
            
# 워드 파일 생성 함수
def create_word_document(keyword, final_selection, analysis=""):
    # 새 워드 문서 생성
    doc = docx.Document()
    
    # 제목 스타일 설정
    title = doc.add_heading(f'PwC 뉴스 분석 보고서: {clean_html_entities(keyword)}', level=0)
    for run in title.runs:
        run.font.color.rgb = RGBColor(208, 74, 2)  # PwC 오렌지 색상
    
    # 분석 요약 추가
    if analysis:
        doc.add_heading('회계법인 관점의 분석 결과', level=1)
        doc.add_paragraph(clean_html_entities(analysis))
    
    # 선별된 주요 뉴스 추가
    doc.add_heading('선별된 주요 뉴스', level=1)
    
    for i, news in enumerate(final_selection):
        p = doc.add_paragraph()
        p.add_run(f"{i+1}. {clean_html_entities(news.get('title', ''))}").bold = True
        
        # 날짜 정보 추가
        date_str = news.get('date', '날짜 정보 없음')
        date_paragraph = doc.add_paragraph()
        date_paragraph.add_run(f"날짜: {clean_html_entities(date_str)}").italic = True
        
        # 선정 사유 추가
        reason = news.get('reason', '')
        if reason:
            doc.add_paragraph(f"선정 사유: {clean_html_entities(reason)}")
        
        # 키워드 추가
        keywords = news.get('keywords', [])
        if keywords:
            doc.add_paragraph(f"키워드: {', '.join([clean_html_entities(k) for k in keywords])}")
        
        # 관련 계열사 추가
        affiliates = news.get('affiliates', [])
        if affiliates:
            doc.add_paragraph(f"관련 계열사: {', '.join([clean_html_entities(a) for a in affiliates])}")
        
        # 언론사 추가
        press = news.get('press', '알 수 없음')
        doc.add_paragraph(f"언론사: {clean_html_entities(press)}")
        
        # URL 추가
        url = news.get('url', '')
        if url:
            doc.add_paragraph(f"출처: {clean_html_entities(url)}")
        
        # 구분선 추가
        if i < len(final_selection) - 1:
            doc.add_paragraph("").add_run().add_break()
    
    # 날짜 및 푸터 추가
    current_date = datetime.now().strftime("%Y년 %m월 %d일")
    doc.add_paragraph(f"\n보고서 생성일: {current_date}")
    doc.add_paragraph("© 2024 PwC 뉴스 분석기 | 회계법인 관점의 뉴스 분석 도구")
    
    return doc

# BytesIO 객체로 워드 문서 저장
def get_binary_file_downloader_html(doc, file_name):
    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

# 커스텀 CSS
st.markdown("""
<style>
    .title-container {
        display: flex;
        align-items: center;
        gap: 20px;
        margin-bottom: 20px;
    }
    .main-title {
        color: #d04a02;
        font-size: 2.5rem;
        font-weight: 700;
    }
    .news-card {
        background-color: #f9f9f9;
        border-radius: 10px;
        padding: 15px;
        margin-bottom: 15px;
        border-left: 4px solid #d04a02;
    }
    .news-title {
        font-weight: 600;
        font-size: 1.1rem;
    }
    .news-url {
        color: #666;
        font-size: 0.9rem;
    }
    .news-date {
        color: #666;
        font-size: 0.9rem;
        font-style: italic;
        margin-top: 5px;
    }
    .analysis-box {
        background-color: #f5f5ff;
        border-radius: 10px;
        padding: 20px;
        margin: 20px 0;
        border-left: 4px solid #d04a02;
    }
    .subtitle {
        color: #dc582a;
        font-size: 1.3rem;
        font-weight: 600;
        margin-top: 20px;
        margin-bottom: 10px;
    }
    .download-box {
        background-color: #eaf7f0;
        border-radius: 10px;
        padding: 20px;
        margin: 20px 0;
        border-left: 4px solid #00a36e;
        text-align: center;
    }
    .analysis-section {
        background-color: #f8f9fa;
        border-left: 4px solid #d04a02;
        padding: 20px;
        margin: 10px 0;
        border-radius: 5px;
    }
    .selected-news {
        border-left: 4px solid #0077b6;
        padding: 15px;
        margin: 10px 0;
        background-color: #f0f8ff;
        border-radius: 5px;
    }
    .excluded-news {
        color: #666;
        padding: 5px 0;
        margin: 5px 0;
        font-size: 0.9em;
    }
    .news-meta {
        color: #666;
        font-size: 0.9em;
        margin: 3px 0;
    }
    .selection-reason {
        color: #666;
        margin: 5px 0;
        font-size: 0.95em;
    }
    .keywords {
        color: #666;
        font-size: 0.9em;
        margin: 5px 0;
    }
    .affiliates {
        color: #666;
        font-size: 0.9em;
        margin: 5px 0;
    }
    .news-url {
        color: #0077b6;
        font-size: 0.9em;
        margin: 5px 0;
        word-break: break-all;
    }
    .news-title-large {
        font-size: 1.2em;
        font-weight: 600;
        color: #000;
        margin-bottom: 8px;
        line-height: 1.4;
    }
    .news-url {
        color: #0077b6;
        font-size: 0.9em;
        margin: 5px 0 10px 0;
        word-break: break-all;
    }
    .news-summary {
        color: #444;
        font-size: 0.95em;
        margin: 10px 0;
        line-height: 1.4;
    }
    .selection-reason {
        color: #666;
        font-size: 0.95em;
        margin: 10px 0;
        line-height: 1.4;
    }
    .importance-high {
        color: #d04a02;
        font-weight: 700;
        margin: 5px 0;
    }
    .importance-medium {
        color: #0077b6;
        font-weight: 700;
        margin: 5px 0;
    }
    .group-indices {
        color: #666;
        font-size: 0.9em;
    }
    .group-selected {
        color: #00a36e;
        font-weight: 600;
    }
    .group-reason {
        color: #666;
        font-size: 0.9em;
        margin-top: 5px;
    }
    .not-selected-news {
        color: #666;
        padding: 5px 0;
        margin: 5px 0;
        font-size: 0.9em;
    }
    .importance-low {
        color: #666;
        font-weight: 700;
        margin: 5px 0;
    }
    .not-selected-reason {
        color: #666;
        margin: 5px 0;
        font-size: 0.95em;
    }
    .email-preview {
        background-color: white;
        border: 1px solid #ddd;
        border-radius: 5px;
        padding: 20px;
        margin: 20px 0;
        overflow-y: auto;
        max-height: 500px;
    }
    .copy-button {
        background-color: #d04a02;
        color: white;
        padding: 10px 20px;
        border: none;
        border-radius: 5px;
        cursor: pointer;
        margin: 10px 0;
    }
    .copy-button:hover {
        background-color: #b33d00;
    }
</style>
""", unsafe_allow_html=True)

# 로고와 제목
col1, col2 = st.columns([1, 5])
with col1:
    # 로고 표시
    logo_path = "logo_orange.png"
    if os.path.exists(logo_path):
        st.image(logo_path, width=100)
    else:
        st.error("로고 파일을 찾을 수 없습니다. 프로젝트 루트에 'logo_orange.png' 파일을 추가해주세요.")

with col2:
    # 메인 타이틀 (로고 포함)
    col1, col2, col3 = st.columns([1, 3, 1])
    
    with col1:
        st.image("logo_orange.png", width=80)
    
    with col2:
        st.markdown("<h1 class='main-title'>PwC 뉴스 분석기</h1>", unsafe_allow_html=True)
    
    with col3:
        st.write("")  # 빈 공간
    
    st.markdown("회계법인 관점에서 중요한 뉴스를 자동으로 분석하는 AI 도구")
    
    # 브라우저 탭 제목 설정
    st.markdown("<script>document.title = 'PwC 뉴스 분석기';</script>", unsafe_allow_html=True)

# 기본 선택 키워드 카테고리를 삼일PwC_핵심으로 설정
    DEFAULT_KEYWORDS = COMPANY_CATEGORIES["Anchor"]

# 사이드바 설정
    st.sidebar.title("🔍 PwC 뉴스 분석기")

# 0단계: 기본 설정
st.sidebar.markdown("### 📋 0단계: 기본 설정")

# 유효 언론사 설정
valid_press_dict = st.sidebar.text_area(
    "📰 유효 언론사 설정 ",
    value="""조선일보: ["조선일보", "chosun", "chosun.com"]
    중앙일보: ["중앙일보", "joongang", "joongang.co.kr", "joins.com"]
    동아일보: ["동아일보", "donga", "donga.com"]
    조선비즈: ["조선비즈", "chosunbiz", "biz.chosun.com"]
    매거진한경: ["매거진한경", "magazine.hankyung", "magazine.hankyung.com"]
    한국경제: ["한국경제", "한경", "hankyung", "hankyung.com", "한경닷컴"]
    매일경제: ["매일경제", "매경", "mk", "mk.co.kr"]
    연합뉴스: ["연합뉴스", "yna", "yna.co.kr"]
    파이낸셜뉴스: ["파이낸셜뉴스", "fnnews", "fnnews.com"]
    데일리팜: ["데일리팜", "dailypharm", "dailypharm.com"]
    IT조선: ["it조선", "it.chosun.com", "itchosun"]
    머니투데이: ["머니투데이", "mt", "mt.co.kr"]
    비즈니스포스트: ["비즈니스포스트", "businesspost", "businesspost.co.kr"]
    이데일리: ["이데일리", "edaily", "edaily.co.kr"]
    아시아경제: ["아시아경제", "asiae", "asiae.co.kr"]
    뉴스핌: ["뉴스핌", "newspim", "newspim.com"]
    뉴시스: ["뉴시스", "newsis", "newsis.com"]
    헤럴드경제: ["헤럴드경제", "herald", "heraldcorp", "heraldcorp.com"]""",
    help="분석에 포함할 신뢰할 수 있는 언론사와 그 별칭을 설정하세요. 형식: '언론사: [별칭1, 별칭2, ...]'",
    key="valid_press_dict"
)

# 추가 언론사 설정 (재평가 시에만 사용됨)
additional_press_dict = st.sidebar.text_area(
    "📰 추가 언론사 설정 (재평가 시에만 사용)",
    value="""철강금속신문: ["철강금속신문", "snmnews", "snmnews.com"]
    에너지신문: ["에너지신문", "energy-news", "energy-news.co.kr"]
    이코노믹데일리: ["이코노믹데일리", "economidaily", "economidaily.com"]""",
    help="기본 언론사에서 뉴스가 선택되지 않을 경우, 재평가 단계에서 추가로 고려할 언론사와 별칭을 설정하세요. 형식: '언론사: [별칭1, 별칭2, ...]'",
    key="additional_press_dict"
)

# 구분선 추가
st.sidebar.markdown("---")

# 날짜 필터 설정
st.sidebar.markdown("### 📅 날짜 필터")

# 현재 시간 가져오기
now = datetime.now()

# 기본 시작 날짜/시간 계산
default_start_date = now - timedelta(days=1)

# Set time to 8:00 AM for both start and end - 한국 시간 기준
start_datetime = datetime.combine(default_start_date.date(), 
                                    datetime.strptime("08:00", "%H:%M").time(), KST)
end_datetime = datetime.combine(now.date(), 
                                datetime.strptime("08:00", "%H:%M").time(), KST)

col1, col2 = st.sidebar.columns(2)
with col1:
    start_date = st.date_input(
        "시작 날짜",
        value=default_start_date.date(),
        help="이 날짜부터 뉴스를 검색합니다. 월요일인 경우 지난 금요일, 그 외에는 전일로 자동 설정됩니다."
    )
    start_time = st.time_input(
        "시작 시간",
        value=start_datetime.time(),
        help="시작 날짜의 구체적인 시간을 설정합니다. 기본값은 오전 8시입니다."
    )
with col2:
    end_date = st.date_input(
        "종료 날짜",
        value=now.date(),
        help="이 날짜까지의 뉴스를 검색합니다."
    )
    end_time = st.time_input(
        "종료 시간",
        value=end_datetime.time(),
        help="종료 날짜의 구체적인 시간을 설정합니다. 기본값은 오전 8시입니다."
    )

# 구분선 추가
st.sidebar.markdown("---")

# 키워드 선택 UI
st.sidebar.markdown("### 🔍 분석할 키워드 선택")

# 키워드 카테고리 선택
selected_category = st.sidebar.radio(
    "키워드 카테고리를 선택하세요",
    options=list(KEYWORD_CATEGORIES.keys()),
    index=0,  # 삼일PwC_핵심을 기본값으로 설정
    help="분석할 키워드 카테고리를 선택하세요. 삼일PwC_핵심 또는 회계업계_일반 중에서 선택할 수 있습니다."
)

# 선택된 카테고리에 따라 키워드 목록 가져오기
SELECTED_KEYWORDS = KEYWORD_CATEGORIES[selected_category]

# 카테고리 내 키워드들 표시
st.sidebar.markdown("**해당 카테고리의 키워드들:**")
for keyword in SELECTED_KEYWORDS:
    st.sidebar.info(f"🔑 {keyword}")

# 선택된 키워드들
selected_keywords = SELECTED_KEYWORDS.copy()

# 선택된 키워드 정보 표시
st.sidebar.markdown("---")
st.sidebar.markdown("### ℹ️ 선택된 키워드 정보")
st.sidebar.info(f"**선택된 키워드:** {len(selected_keywords)}개")

# 키워드별 연관 검색어 정보
st.sidebar.markdown("### 🔍 키워드별 연관 검색어")

# 세션 상태에 COMPANY_KEYWORD_MAP 저장 (초기화)
if 'company_keyword_map' not in st.session_state:
    st.session_state.company_keyword_map = COMPANY_KEYWORD_MAP.copy()

# 선택된 키워드들의 연관 검색어 표시
if selected_keywords:
    for keyword in selected_keywords:
        related_keywords = st.session_state.company_keyword_map.get(keyword, [keyword])
        st.sidebar.info(f"**{keyword}**: {', '.join(related_keywords[:5])}...")

# 미리보기 버튼
with st.sidebar.expander("🔍 검색 키워드 미리보기"):
    if selected_keywords:
        st.info(f"**{len(selected_keywords)}개 키워드가 선택되어 검색됩니다.**")
        for keyword in selected_keywords:
            st.write(f"• {keyword}")
    else:
        st.info("키워드가 선택되지 않았습니다.")

# 검색용 키워드 리스트 (선택된 키워드 + 연관 검색어)
search_keywords = []
for keyword in selected_keywords:
    # 키워드 자체와 연관 검색어 모두 추가
    related_keywords = st.session_state.company_keyword_map.get(keyword, [keyword])
    search_keywords.extend(related_keywords)

# 중복 제거
search_keywords = list(set(search_keywords))

# 구분선 추가
st.sidebar.markdown("---")

# 특화 기준 관리 섹션
st.sidebar.markdown("### 🎯 특화 기준 정보")

# 삼일PwC 키워드인지 확인
is_samil_pwc = any(keyword in ["삼일회계", "삼일PwC", "PwC삼일", "PwC코리아"] for keyword in selected_keywords)

if is_samil_pwc:
    st.sidebar.success("**삼일PwC 특별 프롬프트가 적용됩니다!**")
    st.sidebar.info("삼일PwC 관련 뉴스에 대해 상세한 포함/제외 기준이 적용됩니다.")
else:
    st.sidebar.info("**일반 회계법인 기준이 적용됩니다.**")

# 미리보기 버튼
with st.sidebar.expander("🔍 특화 기준 미리보기"):
    if is_samil_pwc:
        st.success("**삼일PwC 특별 기준 적용**")
        st.write("• 상세한 포함/제외 기준")
        st.write("• 중복 제거 기준")
        st.write("• 경계 사례 판단 기준")
    else:
        st.info("**일반 회계법인 기준 적용**")
        st.write("• 기본 제외/선택 기준")

# 구분선 추가
st.sidebar.markdown("---")

# GPT 모델 선택 섹션
st.sidebar.markdown("### 🤖 GPT 모델 선택")

selected_model = st.sidebar.selectbox(
    "분석에 사용할 GPT 모델을 선택하세요",
    options=list(GPT_MODELS.keys()),
    index=list(GPT_MODELS.keys()).index(DEFAULT_GPT_MODEL) if DEFAULT_GPT_MODEL in GPT_MODELS else 0,
    format_func=lambda x: f"{x} - {GPT_MODELS[x]}",
    help="각 모델의 특성:\n" + "\n".join([f"• {k}: {v}" for k, v in GPT_MODELS.items()])
)

# 모델 설명 표시
st.sidebar.markdown(f"""
<div style='background-color: #f0f2f6; padding: 10px; border-radius: 5px; margin-bottom: 20px;'>
    <strong>선택된 모델:</strong> {selected_model}<br>
    <strong>특징:</strong> {GPT_MODELS[selected_model]}
</div>
""", unsafe_allow_html=True)

# 구분선 추가
st.sidebar.markdown("---")

# 검색 결과 수 - 키워드당 200개로 설정 (신뢰할 수 있는 언론사에서만)
max_results = 200

# AI 프롬프트 설정 (사용자 편집 가능)
st.sidebar.markdown("### 🤖 AI 프롬프트 설정")
st.sidebar.info("AI 분석에 사용되는 프롬프트는 config.py에서 관리됩니다.")

st.sidebar.markdown("---")
st.sidebar.markdown("### 📋 1단계: 제외 판단 기준")

# 제외 기준 설정 - 기본 기준만 표시하고 사용자 수정 허용
exclusion_criteria = st.sidebar.text_area(
    "❌ 제외 기준",
    value=EXCLUSION_CRITERIA,
    help="분석에서 제외할 뉴스의 기준을 설정하세요. 실제 분석 시 각 회사별 특화 기준이 추가로 적용됩니다.",
    key="exclusion_criteria",
    height=300
)


# 구분선 추가
st.sidebar.markdown("---")

# 2단계: 그룹핑 기준
st.sidebar.markdown("### 📋 2단계: 그룹핑 기준")

# 중복 처리 기준 설정 - 기본 기준만 표시하고 사용자 수정 허용
duplicate_handling = st.sidebar.text_area(
    "🔄 중복 처리 기준",
    value=DUPLICATE_HANDLING,
    help="중복된 뉴스를 처리하는 기준을 설정하세요. 실제 분석 시 각 회사별 특화 기준이 추가로 적용됩니다.",
    key="duplicate_handling",
    height=300
)

# 구분선 추가
st.sidebar.markdown("---")

# 3단계: 선택 기준
st.sidebar.markdown("### 📋 3단계: 선택 기준")

# 선택 기준 설정 - 기본 기준만 표시하고 사용자 수정 허용
selection_criteria = st.sidebar.text_area(
    "✅ 선택 기준",
    value=SELECTION_CRITERIA,
    help="뉴스 선택에 적용할 주요 기준들을 나열하세요. 실제 분석 시 각 회사별 특화 기준이 추가로 적용됩니다.",
    key="selection_criteria",
    height=300
)

# 응답 형식 설정
response_format = st.sidebar.text_area(
    "📝 응답 형식",
    value="""선택된 뉴스 인덱스: [1, 3, 5]와 같은 형식으로 알려주세요.

각 선택된 뉴스에 대해:
제목: (뉴스 제목)
언론사: (언론사명)
발행일: (발행일자)
선정 사유: (구체적인 선정 이유)
분석 키워드: (해당 기업 그룹의 주요 계열사들)

[제외된 주요 뉴스]
제외된 중요 뉴스들에 대해:
인덱스: (뉴스 인덱스)
제목: (뉴스 제목)
제외 사유: (구체적인 제외 이유)""",
    help="분석 결과의 출력 형식을 설정하세요.",
    key="response_format",
    height=200
)

# 최종 프롬프트 생성
analysis_prompt = f"""
당신은 회계법인의 전문 애널리스트입니다. 아래 뉴스 목록을 분석하여 회계법인 관점에서 가장 중요한 뉴스를 선별하세요. 

[선택 기준]
{selection_criteria}

[제외 대상]
{exclusion_criteria}

[응답 요구사항]
1. 선택 기준에 부합하는 뉴스가 많다면 최대 3개까지 선택 가능합니다.
2. 선택 기준에 부합하는 뉴스가 없다면, 그 이유를 명확히 설명해주세요.

[응답 형식]
다음과 같은 JSON 형식으로 응답해주세요:

{{
    "selected_news": [
        {{
            "index": 1,
            "title": "뉴스 제목",
            "press": "언론사명",
            "date": "발행일자",
            "reason": "선정 사유",
            "keywords": ["키워드1", "키워드2"]
        }},
        ...
    ],
    "excluded_news": [
        {{
            "index": 2,
            "title": "뉴스 제목",
            "reason": "제외 사유"
        }},
        ...
    ]
}}

[유효 언론사]
{valid_press_dict}

[중복 처리 기준]
{duplicate_handling}
"""

# 메인 컨텐츠
if st.button("뉴스 분석 시작", type="primary"):
    # 뉴스 분석 서비스 초기화
    news_service = NewsAnalysisService()
    
    # 유효 언론사 설정을 딕셔너리로 파싱
    valid_press_config = parse_press_config(valid_press_dict)
    
    # 이메일 미리보기를 위한 전체 내용 저장
    email_content = "[Client Intelligence]\n\n"
    
    # 모든 키워드 분석 결과를 저장할 딕셔너리
    all_results = {}
    
    # 분석 프롬프트 설정
    analysis_prompt = f"""
    당신은 회계법인의 전문 애널리스트입니다. 아래 뉴스 목록을 분석하여 회계법인 관점에서 가장 중요한 뉴스를 선별하세요. 
    
    [선택 기준]
    {selection_criteria}
    
    [제외 대상]
    {exclusion_criteria}
    
    [응답 요구사항]
    1. 선택 기준에 부합하는 뉴스가 많다면 최대 3개까지 선택 가능합니다.
    2. 선택 기준에 부합하는 뉴스가 없다면, 그 이유를 명확히 설명해주세요.
    
    [응답 형식]
    다음과 같은 JSON 형식으로 응답해주세요:
    
    {{
        "selected_news": [
            {{
                "index": 1,
                "title": "뉴스 제목",
                "press": "언론사명",
                "date": "발행일자",
                "reason": "선정 사유",
                "keywords": ["키워드1", "키워드2"]
            }},
            ...
        ],
        "excluded_news": [
            {{
                "index": 2,
                "title": "뉴스 제목",
                "reason": "제외 사유"
            }},
            ...
        ]
    }}
    
    [유효 언론사]
    {valid_press_dict}
    
    [중복 처리 기준]
    {duplicate_handling}
    """
    st.info("📊 **회계법인 기준 적용됨**")
    
    # 키워드별 분석 실행
    for i, keyword in enumerate(selected_keywords, 1):
        with st.spinner(f"'{keyword}' 관련 뉴스를 수집하고 분석 중입니다..."):
            # 해당 키워드의 연관 검색어 확장
            related_keywords = st.session_state.company_keyword_map.get(keyword, [keyword])
            
            # 연관 검색어 표시
            st.write(f"'{keyword}' 연관 검색어로 검색 중: {', '.join(related_keywords)}")
            
            # 날짜/시간 객체 생성
            start_dt = datetime.combine(start_date, start_time)
            end_dt = datetime.combine(end_date, end_time)
            
            # 뉴스 분석 서비스 호출
            try:
                analysis_result = news_service.analyze_news(
                    keywords=related_keywords,
                    start_date=start_dt,
                    end_date=end_dt,
                    companies=[keyword],
                    trusted_press=valid_press_config
                )
                
                # 결과 저장
                all_results[keyword] = analysis_result
                
                # 결과 표시
                st.success(f"'{keyword}' 분석 완료!")
                st.write(f"수집된 뉴스: {analysis_result['collected_count']}개")
                st.write(f"날짜 필터링 후: {analysis_result['date_filtered_count']}개")
                st.write(f"언론사 필터링 후: {analysis_result['press_filtered_count']}개")
                st.write(f"최종 선별: {len(analysis_result['final_selection'])}개")
                
                # 최종 선별된 뉴스 표시
                if analysis_result['final_selection']:
                    st.subheader(f"📰 {keyword} 최종 선별 뉴스")
                    for j, news in enumerate(analysis_result['final_selection'], 1):
                        with st.expander(f"{j}. {news.get('content', '제목 없음')}"):
                            st.write(f"**언론사:** {news.get('press', '알 수 없음')}")
                            st.write(f"**날짜:** {news.get('date', '날짜 정보 없음')}")
                            st.write(f"**URL:** {news.get('url', '')}")
                
            except Exception as e:
                st.error(f"'{keyword}' 분석 중 오류 발생: {str(e)}")
                continue
            
            # 분석 완료 후 결과 요약
            st.success(f"✅ {keyword} 분석 완료!")
            
            # 이메일 내용에 추가
            email_content += f"\n=== {keyword} 분석 결과 ===\n"
            email_content += f"수집된 뉴스: {analysis_result['collected_count']}개\n"
            email_content += f"날짜 필터링 후: {analysis_result['date_filtered_count']}개\n"
            email_content += f"언론사 필터링 후: {analysis_result['press_filtered_count']}개\n"
            email_content += f"최종 선별: {len(analysis_result['final_selection'])}개\n\n"
            
            # 보류 뉴스
            with st.expander("⚠️ 보류 뉴스"):
                for news in analysis_result["borderline_news"]:
                    st.markdown(f"<div class='excluded-news'>[{news['index']}] {news['title']}<br/>└ {news['reason']}</div>", unsafe_allow_html=True)
            
            # 유지 뉴스
            with st.expander("✅ 유지 뉴스"):
                for news in analysis_result["retained_news"]:
                    st.markdown(f"<div class='excluded-news'>[{news['index']}] {news['title']}<br/>└ {news['reason']}</div>", unsafe_allow_html=True)
            
            # 4단계: 그룹핑 결과 표시
            st.markdown("<div class='subtitle'>🔍 4단계: 뉴스 그룹핑 결과</div>", unsafe_allow_html=True)
            
            with st.expander("📋 그룹핑 결과 보기"):
                for group in analysis_result["grouped_news"]:
                    st.markdown(f"""
                    <div class="analysis-section">
                        <h4>그룹 {group['indices']}</h4>
                        <p>선택된 기사: {group['selected_index']}</p>
                        <p>선정 이유: {group['reason']}</p>
                    </div>
                    """, unsafe_allow_html=True)
            
            # 5단계: 최종 선택 결과 표시
            st.markdown("<div class='subtitle'>🔍 5단계: 최종 선택 결과</div>", unsafe_allow_html=True)
            
            # 재평가 여부 확인
            was_reevaluated = analysis_result.get("is_reevaluated", False)
            
            if was_reevaluated:
                st.warning("5단계에서 선정된 뉴스가 없어 6단계 재평가를 진행했습니다.")
                st.markdown("<div class='subtitle'>🔍 6단계: 재평가 결과</div>", unsafe_allow_html=True)
                st.markdown("### 📰 재평가 후 선정된 뉴스")
                news_style = "border-left: 4px solid #FFA500; background-color: #FFF8DC;"
                reason_prefix = "<span style=\"color: #FFA500; font-weight: bold;\">재평가 후</span> 선별 이유: "
            else:
                st.markdown("### 📰 최종 선정된 뉴스")  
                news_style = ""
                reason_prefix = "선별 이유: "
            
            # 최종 선정된 뉴스 표시
            for news in analysis_result["final_selection"]:
                date_str = format_date(news.get('date', ''))
                
                try:
                    date_obj = datetime.strptime(date_str, '%Y-%m-%d')
                    formatted_date = date_obj.strftime('%m/%d')
                except Exception as e:
                    try:
                        date_obj = datetime.strptime(date_str, '%a, %d %b %Y %H:%M:%S %Z')
                        formatted_date = date_obj.strftime('%m/%d')
                    except Exception as e:
                        formatted_date = date_str if date_str else '날짜 정보 없음'

                url = news.get('url', 'URL 정보 없음')
                press = news.get('press', '언론사 정보 없음')
                
                st.markdown(f"""
                    <div class="selected-news" style="{news_style}">
                        <div class="news-title-large">{news['title']} ({formatted_date})</div>
                        <div class="news-url">🔗 <a href="{url}" target="_blank">{url}</a></div>
                        <div class="selection-reason">
                            • {reason_prefix}{news['reason']}
                        </div>
                        <div class="news-summary">
                            • 키워드: {', '.join(news['keywords'])} | 관련 계열사: {', '.join(news['affiliates'])} | 언론사: {press}
                        </div>
                    </div>
                """, unsafe_allow_html=True)
                
                st.markdown("---")
            
            # 디버그 정보
            st.info("AI 분석이 완료되었습니다. 상세한 분석 과정은 로그에서 확인할 수 있습니다.")
            
            # 이메일 내용 추가
            email_content += f"{i}. {keyword}\n"
            for news in analysis_result["final_selection"]:
                date_str = news.get('date', '')
                try:
                    date_obj = datetime.strptime(date_str, '%Y-%m-%d')
                    formatted_date = date_obj.strftime('%m/%d')
                except Exception as e:
                    try:
                        date_obj = datetime.strptime(date_str, '%a, %d %b %Y %H:%M:%S %Z')
                        formatted_date = date_obj.strftime('%m/%d')
                    except Exception as e:
                        formatted_date = date_str if date_str else '날짜 정보 없음'
                
                url = news.get('url', '')
                email_content += f"  - {news['title']} ({formatted_date}) {url}\n"
            email_content += "\n"
            
            st.markdown("---")

    # 모든 키워드 분석이 끝난 후 이메일 미리보기 섹션 추가
    st.markdown("<div class='subtitle'>📧 이메일 미리보기</div>", unsafe_allow_html=True)
    
    # HTML 버전 생성
    html_email_content = "<div style='font-family: Arial, sans-serif; max-width: 800px; font-size: 14px; line-height: 1.5;'>"
    
    html_email_content += "<div style='margin-top: 20px; font-size: 14px;'>안녕하세요, 좋은 아침입니다!<br>오늘의 Client Intelligence 전달 드립니다.<br><br></div>"
    plain_email_content = "\n안녕하세요, 좋은 아침입니다!\n오늘의 Client Intelligence 전달 드립니다."
    
    html_email_content += "<div style='font-size: 14px; font-weight: bold; margin-bottom: 15px; border-bottom: 1px solid #000;'>[Client Intelligence]</div>"
    
    # 일반 텍스트 버전 생성 (복사용)
    plain_email_content += "[Client Intelligence]\n\n"
    
    def clean_title(title):
        """Clean title by removing the press name pattern at the end"""
        # Remove the press pattern (e.g., '제목 - 조선일보', '제목-조선일보', '제목 - Chosun Biz')
        title = re.sub(r"\s*-\s*[가-힣A-Za-z0-9\s]+$", "", title).strip()
        return title

    for i, keyword in enumerate(selected_keywords, 1):
        # HTML 버전에서 키워드를 파란색으로 표시
        html_email_content += f"<div style='font-size: 14px; font-weight: bold; margin-top: 15px; margin-bottom: 10px; color: #0000FF;'>{i}. {keyword}</div>"
        html_email_content += "<ul style='list-style-type: none; padding-left: 20px; margin: 0;'>"
        
        # 텍스트 버전에서도 키워드 구분을 위해 줄바꿈 추가
        plain_email_content += f"{i}. {keyword}\n"
        
        # 해당 키워드의 뉴스 가져오기
        news_list = all_results.get(keyword, [])
        
        if not news_list:
            # 최종 선정 뉴스가 0건인 경우 안내 문구 추가
            html_email_content += "<li style='margin-bottom: 8px; font-size: 14px; color: #888;'>AI 분석결과 금일자로 회계법인 관점에서 특별히 주목할 만한 기사가 없습니다.</li>"
            plain_email_content += "  - AI 분석결과 금일자로 회계법인 관점에서 특별히 주목할 만한 기사가 없습니다.\n"
        else:
            for news in news_list:
                # 날짜 형식 변환
                date_str = news.get('date', '')
                try:
                    date_obj = datetime.strptime(date_str, '%Y-%m-%d')
                    formatted_date = date_obj.strftime('%m/%d')
                except Exception as e:
                    try:
                        date_obj = datetime.strptime(date_str, '%a, %d %b %Y %H:%M:%S %Z')
                        formatted_date = date_obj.strftime('%m/%d')
                    except Exception as e:
                        formatted_date = date_str if date_str else '날짜 정보 없음'
                
                url = news.get('url', '')
                title = news.get('title', '')
                # 이메일 미리보기에서는 언론사 패턴 제거
                title = clean_title(title)
                # HTML 버전 - 링크를 [파일 링크]로 표시하고 글자 크기 통일, 본문 bold 처리
                html_email_content += f"<li style='margin-bottom: 8px; font-size: 14px;'><span style='font-weight: bold;'>- {title} ({formatted_date})</span> <a href='{url}' style='color: #1a0dab; text-decoration: none;'>[기사 링크]</a></li>"
                
                # 텍스트 버전 - 링크를 [파일 링크]로 표시하고 실제 URL은 그 다음 줄에
                plain_email_content += f"  - {title} ({formatted_date}) [기사 링크]\n    {url}\n"
        
        html_email_content += "</ul>"
        plain_email_content += "\n"
    
    # 서명 추가
    html_email_content += "<div style='margin-top: 20px; font-size: 14px;'><br>감사합니다.<br>Client & Market 드림</div>"
    plain_email_content += "\n감사합니다.\nClient & Market 드림"
    
    html_email_content += "</div>"
    
    # 이메일 미리보기 표시
    st.markdown(f"<div class='email-preview'>{html_email_content}</div>", unsafe_allow_html=True)

    # 워드 문서 다운로드 섹션 추가
    st.markdown("<div class='subtitle'>📄 워드 문서 다운로드</div>", unsafe_allow_html=True)
    
    # 워드 문서 생성
    if all_results:
        try:
            # 모든 키워드의 최종 선별 뉴스를 하나의 워드 문서로 생성
            all_final_news = []
            for keyword, result in all_results.items():
                if 'final_selection' in result:
                    for news in result['final_selection']:
                        news_with_keyword = news.copy()
                        news_with_keyword['keyword'] = keyword
                        all_final_news.append(news_with_keyword)
            
            if all_final_news:
                # 워드 문서 생성
                doc = create_word_document("종합 뉴스 분석", all_final_news)
                
                # 현재 날짜로 파일명 생성
                current_date = datetime.now().strftime("%Y%m%d")
                filename = f"PwC_뉴스분석_{current_date}.docx"
                
                # 워드 문서 다운로드 버튼
                bio = get_binary_file_downloader_html(doc, filename)
                st.download_button(
                    label="📥 워드 문서 다운로드 (.docx)",
                    data=bio.getvalue(),
                    file_name=filename,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
                
                st.success("워드 문서가 생성되었습니다. 위 버튼을 클릭하여 다운로드하세요.")
            else:
                st.warning("다운로드할 뉴스가 없습니다.")
                
        except Exception as e:
            st.error(f"워드 문서 생성 중 오류가 발생했습니다: {str(e)}")
    
    # CSV 다운로드 섹션 (인코딩 문제 해결)
    st.markdown("<div class='subtitle'>📊 CSV 다운로드 (인코딩 문제 해결)</div>", unsafe_allow_html=True)
    
    if all_results:
        try:
            # CSV용 데이터 준비
            csv_data = []
            for keyword, result in all_results.items():
                if 'final_selection' in result:
                    for news in result['final_selection']:
                        csv_data.append({
                            '키워드': clean_html_entities(keyword),
                            '제목': clean_html_entities(news.get('title', '')),
                            '날짜': clean_html_entities(news.get('date', '')),
                            '언론사': clean_html_entities(news.get('press', '')),
                            '선별이유': clean_html_entities(news.get('reason', '')),
                            '키워드': clean_html_entities(', '.join(news.get('keywords', []))),
                            '관련계열사': clean_html_entities(', '.join(news.get('affiliates', []))),
                            'URL': clean_html_entities(news.get('url', ''))
                        })
            
            if csv_data:
                # DataFrame 생성
                df = pd.DataFrame(csv_data)
                
                # CSV 파일 생성 (인코딩 문제 해결)
                csv_buffer = io.StringIO()
                # UTF-8 BOM을 추가하여 Excel에서 한글이 깨지지 않도록 함
                df.to_csv(csv_buffer, index=False, encoding='utf-8-sig')
                csv_content = csv_buffer.getvalue()
                
                # CSV 다운로드 버튼
                current_date = datetime.now().strftime("%Y%m%d")
                csv_filename = f"PwC_뉴스분석_{current_date}.csv"
                
                st.download_button(
                    label="📥 CSV 다운로드 (.csv) - 인코딩 문제 해결됨",
                    data=csv_content,
                    file_name=csv_filename,
                    mime="text/csv"
                )
                
                st.success("CSV 파일이 생성되었습니다. UTF-8 BOM 인코딩으로 한글이 깨지지 않습니다.")
                
                # 미리보기 표시 (클릭 가능한 링크 포함)
                st.markdown("**CSV 미리보기:**")
                
                # HTML 테이블로 표시하여 링크를 클릭 가능하게 만듦
                html_table = df.to_html(
                    index=False, 
                    escape=False,  # HTML 이스케이프 방지
                    formatters={
                        'URL': lambda x: f'<a href="{x}" target="_blank" style="color: #0077b6; text-decoration: underline;">🔗 링크</a>' if x else ''
                    }
                )
                
                # HTML 테이블 스타일링
                styled_html = f"""
                <div style="overflow-x: auto;">
                    <style>
                        table {{
                            border-collapse: collapse;
                            width: 100%;
                            font-family: Arial, sans-serif;
                        }}
                        th, td {{
                            border: 1px solid #ddd;
                            padding: 8px;
                            text-align: left;
                            vertical-align: top;
                        }}
                        th {{
                            background-color: #f2f2f2;
                            font-weight: bold;
                        }}
                        tr:nth-child(even) {{
                            background-color: #f9f9f9;
                        }}
                        tr:hover {{
                            background-color: #f5f5f5;
                        }}
                        a {{
                            color: #0077b6;
                            text-decoration: underline;
                        }}
                        a:hover {{
                            color: #0056b3;
                        }}
                    </style>
                    {html_table}
                </div>
                """
                
                st.markdown(styled_html, unsafe_allow_html=True)
                
                # 원본 DataFrame도 표시 (데이터 확인용)
                st.markdown("**원본 데이터 (편집용):**")
                st.dataframe(df)
            else:
                st.warning("CSV로 다운로드할 뉴스가 없습니다.")
                
        except Exception as e:
            st.error(f"CSV 생성 중 오류가 발생했습니다: {str(e)}")



else:
    # 초기 화면 설명 (주석 처리됨)
    """
    ### 👋 PwC 뉴스 분석기에 오신 것을 환영합니다!
    
    이 도구는 입력한 키워드에 대한 최신 뉴스를 자동으로 수집하고, 회계법인 관점에서 중요한 뉴스를 선별하여 분석해드립니다.
    
    #### 주요 기능:
    1. 최신 뉴스 자동 수집 (기본 100개)
    2. 신뢰할 수 있는 언론사 필터링
    3. 6단계 AI 기반 뉴스 분석 프로세스:
       - 1단계: 뉴스 수집 - 키워드 기반으로 최신 뉴스 데이터 수집
       - 2단계: 유효 언론사 필터링 - 신뢰할 수 있는 언론사 선별
       - 3단계: 제외/보류/유지 판단 - 회계법인 관점에서의 중요도 1차 분류
       - 4단계: 유사 뉴스 그룹핑 - 중복 기사 제거 및 대표 기사 선정
       - 5단계: 중요도 평가 및 최종 선정 - 회계법인 관점의 중요도 평가
       - 6단계: 필요시 재평가 - 선정된 뉴스가 없을 경우 AI가 기준을 완화하여 재평가
    4. 선별된 뉴스에 대한 상세 정보 제공
       - 제목 및 날짜
       - 원문 링크
       - 선별 이유
       - 키워드, 관련 계열사, 언론사 정보
    5. 분석 결과 이메일 형식 미리보기
    
    #### 사용 방법:
    1. 사이드바에서 분석할 기업을 선택하세요 (최대 10개)
       - 기본 제공 기업 목록에서 선택
       - 새로운 기업 직접 추가 가능
    2. GPT 모델을 선택하세요
       - gpt-4o: 빠르고 실시간 (기본값)
    3. 날짜 필터를 설정하세요
       - 기본값: 어제 또는 지난 금요일(월요일인 경우)부터 오늘까지
    4. "뉴스 분석 시작" 버튼을 클릭하세요
    
    #### 분석 결과 확인:
    - 각 키워드별 최종 선정된 중요 뉴스
    - 선정 과정의 중간 결과(제외/보류/유지, 그룹핑 등)
    - 선정된 모든 뉴스의 요약 이메일 미리보기
    - 디버그 정보 (시스템 프롬프트, AI 응답 등)
    
    """

# 푸터
st.markdown("---")
st.markdown("© 2024 PwC 뉴스 분석기 | 회계법인 관점의 뉴스 분석 도구")
